from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建新文档
doc = Document()

# 添加标题
title = doc.add_heading('"InvestIQ AI"智能金融平台 5周开发计划（优化版）', 0)

# 项目概述
doc.add_heading('项目概述', level=1)
doc.add_paragraph('项目类型: 学校课程项目')
doc.add_paragraph('开发周期: 5周（严格按周推进）')
doc.add_paragraph('当前进度: 完成约30% - 基础框架已搭建，用户认证、股票展示、后端路由、K线图已实现')
doc.add_paragraph('技术选型: 新浪财经API + 实时数据（每日更新） + 多AI模型支持（DeepSeek/GPT/文心/通义等）')

# 团队分工（优化后）
doc.add_heading('团队分工（优化版）', level=1)
doc.add_paragraph('【后端组长】: 刘伯禹 - 负责数据库、股票数据爬取、缓存、核心后端功能')
doc.add_paragraph('【AI模型组】: 程文涛 - 负责大模型API调用、AI适配器开发、提示词管理')
doc.add_paragraph('【用户功能组】: 胡邵旸 - 负责用户偏好设置、收藏夹功能、个性化配置')
doc.add_paragraph('【前端组】: 胡邵阳 + 程文涛 - 共同负责前端UI设计、数据展示、接口对接')

# 第一周
doc.add_heading('第一周：数据库重建 + 基础数据接口准备', level=1)

doc.add_heading('1.1 刘伯禹（后端+数据库）', level=2)
doc.add_paragraph('任务1: 基于现有文档重建数据库')
doc.add_paragraph('• 重新设计数据库结构（9号后的版本）')
doc.add_paragraph('• 编写数据库初始化SQL脚本')
doc.add_paragraph('• 配置MyBatis Mapper对应新表结构')
doc.add_paragraph('交付物: finance_db.sql, 各实体类 Mapper.xml')

doc.add_paragraph('任务2: 新浪财经API集成与数据上传清洗')
doc.add_paragraph('• 对接新浪财经API接口（股票实时/历史数据）')
doc.add_paragraph('• 编写数据上传清洗器 SinaStockDataCleaner')
doc.add_paragraph('• 配置定时任务机制（Spring Scheduler）')
doc.add_paragraph('交付物: SinaApiService.java, StockDataCleanService.java')

doc.add_paragraph('任务3: 股票历史数据表结构设计')
doc.add_paragraph('• 设计 stock_history 表存储历史K线数据')
doc.add_paragraph('• 实现历史数据的增删改查 StockHistoryService')
doc.add_paragraph('交付物: stock_history 表, StockHistoryMapper.xml')

doc.add_heading('1.2 胡邵阳+程文涛（前端UI）', level=2)
doc.add_paragraph('任务1: 股票搜索功能优化')
doc.add_paragraph('• 优化搜索框UI，支持拼音/代码搜索')
doc.add_paragraph('• 添加搜索历史记录显示')
doc.add_paragraph('交付物: 优化后的 SearchBar.vue')

doc.add_paragraph('任务2: 收藏夹页面UI设计')
doc.add_paragraph('• 设计 Watchlist.vue 静态UI')
doc.add_paragraph('• 实现拖拽排序功能（VueDraggable）')
doc.add_paragraph('• 添加股票卡片组件 StockCard.vue')
doc.add_paragraph('交付物: Watchlist.vue, StockCard.vue')

doc.add_paragraph('任务3: 市场行情页面数据展示')
doc.add_paragraph('• 对接后端股票列表API')
doc.add_paragraph('• 实现分页功能')
doc.add_paragraph('• 添加涨跌幅排序')
doc.add_paragraph('交付物: 完善后的 Market.vue')

doc.add_heading('1.3 胡邵旸（前端辅助+文档）', level=2)
doc.add_paragraph('任务1: 前端API模块化封装')
doc.add_paragraph('• 创建统一的API调用文件 api/stock.js, api/advice.js')
doc.add_paragraph('• 配置Axios拦截器（Token自动添加）')
doc.add_paragraph('交付物: api/ 目录完整结构')

doc.add_paragraph('任务2: 编写数据库设计文档')
doc.add_paragraph('• 绘制9张表的ER图说明')
doc.add_paragraph('• 编写各字段说明文档')
doc.add_paragraph('交付物: 数据库设计文档.md')

doc.add_paragraph('任务3: 基础测试框架')
doc.add_paragraph('• 配置Vitest测试环境')
doc.add_paragraph('• 编写第一个测试用例（登录功能）')
doc.add_paragraph('交付物: vitest.config.js, Login.spec.js')

doc.add_paragraph('周末协作: 全员测试数据库连接，确保新表结构正常运行')

# 第二周
doc.add_heading('第二周：基础数据调用 + 核心功能实现', level=1)

doc.add_heading('2.1 刘伯禹（后端+数据）', level=2)
doc.add_paragraph('任务1: 定时任务实现股票数据更新')
doc.add_paragraph('• 使用 @Scheduled 实现每日定时抓取新浪财经数据')
doc.add_paragraph('• 编写增量更新方法 StockDataUpdateScheduler')
doc.add_paragraph('• 实现智能增量更新（只更新变动数据）')
doc.add_paragraph('交付物: StockDataUpdateScheduler.java')

doc.add_paragraph('任务3: 股票数据过滤后端服务')
doc.add_paragraph('• 实现股票数据查询优化API')
doc.add_paragraph('• 添加股票数据缓存机制')
doc.add_paragraph('• 配置股票历史数据过滤')
doc.add_paragraph('交付物: StockDataController.java, CacheService.java')

doc.add_heading('2.2 胡邵旸（用户偏好功能）', level=2)
doc.add_paragraph('任务1: 收藏夹后端服务【核心任务】')
doc.add_paragraph('• 实现收藏夹增删改查API')
doc.add_paragraph('• 设计用户与股票关联表 user_favorite 表结构')
doc.add_paragraph('• 限制收藏夹数量上限（最多50只）')
doc.add_paragraph('交付物: WatchlistController.java, WatchlistService.java')

doc.add_paragraph('任务2: 用户偏好设置后端')
doc.add_paragraph('• 实现用户投资偏好保存API')
doc.add_paragraph('• 设计 user_preference 表')
doc.add_paragraph('• 支持风险偏好、行业偏好等配置')
doc.add_paragraph('交付物: UserPreferenceController.java, UserPreferenceService.java')

doc.add_heading('2.3 胡邵阳+程文涛（前端功能）', level=2)
doc.add_paragraph('任务1: 收藏夹交互功能对接')
doc.add_paragraph('• 对接收藏夹增删改查API')
doc.add_paragraph('• 实现添加/删除股票动画')
doc.add_paragraph('• 添加实时价格刷新（轮询）')
doc.add_paragraph('交付物: 完整功能的 Watchlist.vue')

doc.add_paragraph('任务2: 投资偏好设置页面（与胡邵旸对接）')
doc.add_paragraph('• 创建投资偏好表单页面')
doc.add_paragraph('• 实现风险偏好选择UI')
doc.add_paragraph('• 对接后端偏好保存API')
doc.add_paragraph('交付物: PreferenceSettings.vue')

doc.add_paragraph('任务3: 股票详情页ECharts图表集成')
doc.add_paragraph('• 对接历史数据API')
doc.add_paragraph('• 实现K线图、成交量图、技术指标图')
doc.add_paragraph('• 添加时间范围切换（日/周/月）')
doc.add_paragraph('交付物: 完整图表的 StockDetail.vue')

doc.add_heading('2.4 全员（前端+测试）', level=2)
doc.add_paragraph('任务1: 状态管理配置')
doc.add_paragraph('• 创建 stores/stock.js（股票数据状态）')
doc.add_paragraph('• 创建 stores/watchlist.js（收藏夹状态）')
doc.add_paragraph('• 实现Pinia持久化配置')
doc.add_paragraph('交付物: stores/stock.js, stores/watchlist.js')

doc.add_paragraph('任务2: 核心功能测试编写')
doc.add_paragraph('• 编写收藏夹功能测试')
doc.add_paragraph('• 编写股票搜索测试')
doc.add_paragraph('交付物: Watchlist.spec.js, Market.spec.js')

doc.add_paragraph('任务3: 编写API接口文档')
doc.add_paragraph('• 整理所有后端API接口')
doc.add_paragraph('• 使用Markdown编写接口文档')
doc.add_paragraph('交付物: API接口文档.md')

doc.add_paragraph('周末协作: 全员测试定时任务，验证数据的更新准确性')

# 第三周
doc.add_heading('第三周：AI功能开发 + 核心功能完善', level=1)

doc.add_heading('3.1 程文涛（AI模型集成）【核心任务】', level=2)
doc.add_paragraph('任务1: 多模型适配器开发')
doc.add_paragraph('• 设计统一的AI服务适配器接口 AIServiceAdapter')
doc.add_paragraph('• 实现多个AI模型的适配器类：')
doc.add_paragraph('  - DeepSeekAdapter.java - DeepSeek API')
doc.add_paragraph('  - OpenAIAdapter.java - GPT API')
doc.add_paragraph('  - BaiduAdapter.java - 文心一言API（可选）')
doc.add_paragraph('  - AlibabaAdapter.java - 通义千问API（可选）')
doc.add_paragraph('• 用户可在设置页面配置：API Key + API Endpoint URL')
doc.add_paragraph('• 实现AI请求封装类 AIRequest, AIResponse')
doc.add_paragraph('交付物: AIServiceAdapter 接口, 各个Adapter实现类, AIModelManager.java')

doc.add_paragraph('任务2: AI提示词管理服务')
doc.add_paragraph('• 实现提示词增删改查API')
doc.add_paragraph('• 设计 ai_prompt 表')
doc.add_paragraph('• 配置默认提示词模板（股票推荐、风险分析）')
doc.add_paragraph('交付物: AIPromptController.java, PromptService.java')

doc.add_heading('3.2 刘伯禹（后端辅助）', level=2)
doc.add_paragraph('任务1: 智能投资建议服务（配合程文涛）')
doc.add_paragraph('• 实现基于用户偏好+收藏夹的AI推荐')
doc.add_paragraph('• 编写投资建议生成服务 AdviceService')
doc.add_paragraph('• 设计 investment_advice 表')
doc.add_paragraph('交付物: AdviceController.java, AdviceService.java')

doc.add_heading('3.3 胡邵阳+程文涛（前端AI页面）', level=2)
doc.add_paragraph('任务1: AI模型配置与提示词管理页面')
doc.add_paragraph('• 创建AI模型配置页面 AIModelConfig.vue')
doc.add_paragraph('• 支持添加多个AI模型')
doc.add_paragraph('• 配置API Key和Endpoint URL')
doc.add_paragraph('• 选择默认使用的模型')
doc.add_paragraph('• 创建提示词列表页面 PromptList.vue')
doc.add_paragraph('• 实现提示词创建/编辑功能')
doc.add_paragraph('• 添加提示词测试功能UI（可选择不同模型）')
doc.add_paragraph('交付物: AIModelConfig.vue, PromptList.vue, PromptEditor.vue')

doc.add_paragraph('任务2: 智能投顾页面核心UI')
doc.add_paragraph('• 创建投资建议展示页面 Advisor.vue')
doc.add_paragraph('• 实现"一键生成建议"按钮')
doc.add_paragraph('• 展示AI推荐的股票列表和理由')
doc.add_paragraph('交付物: Advisor.vue')

doc.add_paragraph('任务3: 用户设置页面（含AI配置）')
doc.add_paragraph('• 创建用户设置页面 Settings.vue')
doc.add_paragraph('• 实现密码修改、偏好调整UI')
doc.add_paragraph('• 集成: AI模型管理模块')
doc.add_paragraph('  - 我的AI模型列表')
doc.add_paragraph('  - 添加/编辑/删除AI模型配置')
doc.add_paragraph('  - 设置默认模型')
doc.add_paragraph('交付物: Settings.vue')

doc.add_heading('3.4 胡邵旸（测试+文档）', level=2)
doc.add_paragraph('任务1: 用户手册撰写')
doc.add_paragraph('• 编写平台使用说明')
doc.add_paragraph('• 配图讲解核心功能')
doc.add_paragraph('交付物: 用户手册.md')

doc.add_paragraph('任务2: 前端组件文档')
doc.add_paragraph('• 整理所有Vue组件说明')
doc.add_paragraph('• 编写组件使用指南')
doc.add_paragraph('交付物: 前端组件文档.md')

doc.add_paragraph('周末协作: 全员测试AI功能，验证推荐结果准确性')

# 第四周
doc.add_heading('第四周：回测功能 + 全面测试优化', level=1)

doc.add_heading('4.1 刘伯禹+程文涛（后端+回测）', level=2)
doc.add_paragraph('任务1: 回测引擎框架搭建（含自动优化迭代）')
doc.add_paragraph('• 编写回测引擎 BacktestEngine.java')
doc.add_paragraph('• 实现交易模拟逻辑')
doc.add_paragraph('• 计算收益率、风险指标（年化收益率、最大回撤等）')
doc.add_paragraph('• 核心功能: 回测后自动评估提示词优化')
doc.add_paragraph('  - 回测完成后自动判断结果是否达标')
doc.add_paragraph('  - 根据用户设定阈值，触发提示词优化流程')
doc.add_paragraph('  - 记录优化前后的对比数据')
doc.add_paragraph('交付物: BacktestEngine.java, BacktestService.java, BacktestTrigger.java')

doc.add_paragraph('任务2: 提示词自动优化引擎')
doc.add_paragraph('• 编写提示词优化器 PromptOptimizationEngine.java')
doc.add_paragraph('• 实现AI自动迭代优化生成提示词')
doc.add_paragraph('• 配置迭代次数，基于回测结果调优')
doc.add_paragraph('• 记录迭代日志到 prompt_iteration_log 表')
doc.add_paragraph('交付物: PromptOptimizationEngine.java, 相关数据库SQL')

doc.add_paragraph('任务3: API性能优化和安全加固')
doc.add_paragraph('• 添加API限流（每用户每天最多10次）')
doc.add_paragraph('• 优化SQL查询性能，添加索引')
doc.add_paragraph('• 配置JWT Token刷新机制')
doc.add_paragraph('交付物: 优化后的Controller和配置文件')

doc.add_heading('4.2 胡邵阳+程文涛（前端测试）', level=2)
doc.add_paragraph('任务1: 智能投顾页面完整对接')
doc.add_paragraph('• 对接AI建议生成API')
doc.add_paragraph('• 实现建议结果展示')
doc.add_paragraph('• 添加建议历史记录功能')
doc.add_paragraph('交付物: 完整功能的 Advisor.vue')

doc.add_paragraph('任务2: 回测结果可视化+优化展示')
doc.add_paragraph('• 创建回测结果展示组件 BacktestChart.vue')
doc.add_paragraph('• 使用ECharts展示收益曲线')
doc.add_paragraph('• 展示风险指标卡片')
doc.add_paragraph('• 核心功能: 提示词优化迭代过程可视化')
doc.add_paragraph('  - 展示每次迭代的结果对比')
doc.add_paragraph('  - 显示提示词版本')
doc.add_paragraph('  - 展示优化前后的建议差异')
doc.add_paragraph('交付物: BacktestChart.vue, PromptIterationTimeline.vue')

doc.add_paragraph('任务3: 全局交互优化')
doc.add_paragraph('• 添加Loading动画')
doc.add_paragraph('• 优化错误提示UI')
doc.add_paragraph('• 实现消息通知组件')
doc.add_paragraph('交付物: Loading.vue, Notification.vue')

doc.add_heading('4.3 胡邵旸（测试+部署）', level=2)
doc.add_paragraph('任务1: 端到端测试')
doc.add_paragraph('• 编写完整用户流程测试')
doc.add_paragraph('• 测试从登录到生成建议的完整路径')
doc.add_paragraph('交付物: e2e.spec.js')

doc.add_paragraph('任务2: 性能测试报告')
doc.add_paragraph('• 测试API响应时间')
doc.add_paragraph('• 测试前端页面加载速度')
doc.add_paragraph('交付物: 性能测试报告.md')

doc.add_paragraph('任务3: Bug修复和功能微调')
doc.add_paragraph('• 收集并修复测试发现的Bug')
doc.add_paragraph('• 调整交互细节')

doc.add_paragraph('周末协作: 全员集中修复，清理已知Bug，确保核心流程')

# 第五周
doc.add_heading('第五周：最终部署 + 答辩准备', level=1)

doc.add_heading('5.1 刘伯禹（后端+部署）', level=2)
doc.add_paragraph('任务1: 生产环境部署')
doc.add_paragraph('• 配置生产数据库')
doc.add_paragraph('• 部署后端服务到服务器')
doc.add_paragraph('• 配置Nginx反向代理')
doc.add_paragraph('交付物: 可访问的线上环境')

doc.add_paragraph('任务2: 系统演示准备')
doc.add_paragraph('• 准备演示数据（真实股票数据）')
doc.add_paragraph('• 准备演示账号')
doc.add_paragraph('• 测试演示流程')
doc.add_paragraph('交付物: 演示脚本')

doc.add_paragraph('任务3: 答辩PPT技术部分')
doc.add_paragraph('• 编写技术架构介绍PPT')
doc.add_paragraph('• 准备技术难点讲解')
doc.add_paragraph('• 准备技术答辩讲稿')
doc.add_paragraph('交付物: PPT技术部分（15页）')

doc.add_heading('5.2 胡邵阳+程文涛（前端UI+部署）', level=2)
doc.add_paragraph('任务1: UI最终优化')
doc.add_paragraph('• 统一视觉风格')
doc.add_paragraph('• 优化响应式布局')
doc.add_paragraph('• 修改UI细节问题')
doc.add_paragraph('交付物: 完整版前端代码')

doc.add_paragraph('任务2: 用户体验测试')
doc.add_paragraph('• 邀请同学试用并收集反馈')
doc.add_paragraph('• 优化操作流程')
doc.add_paragraph('交付物: 用户体验报告')

doc.add_paragraph('任务3: 答辩PPT产品演示部分')
doc.add_paragraph('• 录制功能演示视频')
doc.add_paragraph('• 准备产品亮点讲解')
doc.add_paragraph('• 完善PPT页面展示')
doc.add_paragraph('交付物: PPT产品部分（10页）')

doc.add_heading('5.3 胡邵旸（文档+部署）', level=2)
doc.add_paragraph('任务1: 整理项目文档包')
doc.add_paragraph('• 汇总所有技术文档')
doc.add_paragraph('• 编写项目总结报告')
doc.add_paragraph('• 整理代码注释')
doc.add_paragraph('交付物: 完整文档包')

doc.add_paragraph('任务2: 答辩PPT整体整合')
doc.add_paragraph('• 整合完整PPT内容')
doc.add_paragraph('• 统一PPT风格')
doc.add_paragraph('• 编写答辩内容讲稿')
doc.add_paragraph('交付物: 完整答辩PPT（30页）')

doc.add_paragraph('任务3: 答辩预演和QA准备')
doc.add_paragraph('• 组织答辩预演')
doc.add_paragraph('• 准备常见问题回答')
doc.add_paragraph('• 准备应急预案')
doc.add_paragraph('交付物: 答辩QA文档')

doc.add_paragraph('周末协作: 全员答辩预演，模拟老师提问，优化答辩流程')

# 协作机制
doc.add_heading('协作机制', level=1)

doc.add_heading('日常协作', level=2)
doc.add_paragraph('每天9点: 微信群同步进度（群内语音/文字）')
doc.add_paragraph('每周三: 腾讯会议30分钟（周中检查）')
doc.add_paragraph('每周日晚: 周总结/下周计划会（2小时）')

doc.add_heading('代码管理', level=2)
doc.add_paragraph('Git分支策略:')
doc.add_paragraph('• main - 稳定版本')
doc.add_paragraph('• dev - 开发主分支')
doc.add_paragraph('• feature/功能名 - 功能分支')
doc.add_paragraph('提交规范: [类型] 简短描述 [feat] 添加收藏夹功能')
doc.add_paragraph('Code Review: 重要功能需至少一人Review后合并')

doc.add_heading('沟通机制', level=2)
doc.add_paragraph('技术讨论: 微信群内讨论，重要决策会议确定')
doc.add_paragraph('接口对接: 胡邵阳先用Mock数据开发，后端完成后联调')
doc.add_paragraph('进度追踪: 周末从晚上1-2小时追踪进度')

doc.add_heading('关键优化点（重点）', level=1)
doc.add_paragraph('1. 任务分配更清晰：')
doc.add_paragraph('   • 胡邵旸专注用户偏好和收藏夹功能')
doc.add_paragraph('   • 程文涛专注AI模型集成和适配器开发')
doc.add_paragraph('   • 刘伯禹负责核心后端和数据库')
doc.add_paragraph('   • 前端由胡邵阳和程文涛共同完成')

doc.add_paragraph('2. 开发顺序优化：')
doc.add_paragraph('   • 第一周：数据库和基础数据接口（前置依赖）')
doc.add_paragraph('   • 第二周：核心功能（收藏夹、用户偏好）')
doc.add_paragraph('   • 第三周：AI功能开发（依赖前两周的数据）')
doc.add_paragraph('   • 第四周：回测功能（依赖AI和数据完善）')
doc.add_paragraph('   • 第五周：部署和答辩准备')

doc.add_paragraph('3. 回测功能后置原因：')
doc.add_paragraph('   • 需要完整的股票历史数据（第一周）')
doc.add_paragraph('   • 需要AI模型正常工作（第三周）')
doc.add_paragraph('   • 需要前后端接口稳定（第二周）')
doc.add_paragraph('   • 回测是验证性功能，不影响核心流程')

# 保存文档
doc.save('周开发计划_优化版.docx')
print('优化版开发计划已生成！')
