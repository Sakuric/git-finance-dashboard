from docx import Document

doc = Document('周开发计划_优化版.docx')

# 找到并替换团队分工部分
for i, para in enumerate(doc.paragraphs):
    if '团队分工' in para.text:
        # 找到团队分工后的几段进行替换
        if i+1 < len(doc.paragraphs) and '后端+数据库组长' in doc.paragraphs[i+1].text:
            doc.paragraphs[i+1].text = '【后端组长】: 刘伯禹 - 负责数据库、股票数据、缓存、核心后端功能'
        if i+2 < len(doc.paragraphs) and '前端组' in doc.paragraphs[i+2].text:
            doc.paragraphs[i+2].text = '【AI模型组】: 程文涛 - 负责大模型API调用、AI适配器、提示词管理'
        if i+3 < len(doc.paragraphs) and '测试组' in doc.paragraphs[i+3].text:
            doc.paragraphs[i+3].text = '【用户功能组】: 胡邵旸 - 负责用户偏好、收藏夹、个性化设置'
            # 添加前端组
            new_para = doc.paragraphs[i+3].insert_paragraph_before('【前端组】: 胡邵阳 + 程文涛 - 共同负责前端UI、数据展示、接口对接')
        break

# 修改第二周任务2 - 用户偏好功能分配给胡邵旸
for para in doc.paragraphs:
    if '任务2: 收藏夹后端服务' in para.text:
        para.text = '任务2: 收藏夹后端服务【负责人: 胡邵旸】'
    elif '实现收藏夹增删改查API' in para.text:
        next_idx = doc.paragraphs.index(para)
        # 在这段后面添加负责人标注

# 修改第三周任务1 - 大模型API调用分配给程文涛
for para in doc.paragraphs:
    if '任务1: 回测引擎框架搭建' in para.text and '3.1' in doc.paragraphs[doc.paragraphs.index(para)-1].text:
        para.text = '任务1: 多模型适配器开发【负责人: 程文涛】'

doc.save('周开发计划_优化版.docx')
print('文档更新完成')
